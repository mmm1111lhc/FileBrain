"""Word (.docx) 内容提取器"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class WordExtractor:
    """Word 文档内容提取"""

    def __init__(self):
        self._docx = None

    def _lazy_import(self):
        if self._docx is not None:
            return
        try:
            from docx import Document
            self._docx = Document
        except ImportError:
            logger.warning("python-docx 未安装")
            self._docx = False

    def extract(self, file_path: str) -> dict:
        """提取 Word 内容"""
        self._lazy_import()
        path = Path(file_path)
        if not path.exists():
            return {"text": "", "title": "", "success": False,
                    "error": "文件不存在"}

        if not self._docx:
            return {"text": "", "title": "", "success": False,
                    "error": "python-docx 未安装"}

        try:
            doc = self._docx(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)

            # 提取标题：取第一个 Heading 或第一个有意义的段落
            title = ""
            for p in doc.paragraphs:
                if p.style.name.startswith("Heading") and p.text.strip():
                    title = p.text.strip()
                    break
            if not title:
                for p in doc.paragraphs:
                    t = p.text.strip()
                    if 5 < len(t) < 80:
                        title = t
                        break

            # 提取作者
            author = ""
            try:
                author = doc.core_properties.author or ""
            except Exception:
                pass

            return {
                "text": full_text,
                "title": title or "",
                "author": author,
                "success": True,
            }

        except Exception as e:
            logger.error(f"Word 提取失败: {e}")
            return {"text": "", "title": "", "success": False,
                    "error": str(e)}
